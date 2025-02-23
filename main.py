import os
import ttkbootstrap as ttk
import tkinter as tk
from tkinter import filedialog, messagebox, StringVar
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import simpleSplit
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT



# Dictionary of legal document templates
document_templates = {
    "nda": {
        "title": "Non-Disclosure Agreement (NDA)",
        "description": "A contract to protect confidential information shared between two parties.",
        "fields": ["Disclosing Party Name", "Receiving Party Name", "Effective Date"],
        "content": """This Non-Disclosure Agreement ('Agreement') is made as of {Effective Date} by and between:
        
{Disclosing Party Name} (Disclosing Party), and {Receiving Party Name} (Receiving Party).

1. Confidential Information:
   The Disclosing Party agrees to share certain confidential and proprietary information.

2. Obligations of Receiving Party:
   The Receiving Party agrees to keep the information confidential.

3. Term & Termination:
   This Agreement remains in effect until terminated by either party.

"""
    },
    "lease": {
        "title": "Lease Agreement",
        "description": "A legal contract between a landlord and a tenant for rental property.",
        "fields": ["Landlord Name", "Tenant Name", "Property Address", "Rent Amount", "Payment Due Date", "Start Date", "End Date"],
        "content": """This Lease Agreement ('Agreement') is made as of today between:
        
{Landlord Name} (Landlord) and {Tenant Name} (Tenant).

1. Property Description:
   The Landlord leases the property located at {Property Address} to the Tenant.

2. Lease Term:
   The lease starts on {Start Date} and continues until {End Date}.

3. Rent Payment:
   Tenant agrees to pay {Rent Amount} on {Payment Due Date} each month.

4. Governing Law:
   This Agreement is governed by the laws of the relevant jurisdiction.

"""
    },
    "employment": {
        "title": "Employment Contract",
        "description": "A contract defining terms of employment between an employer and an employee.",
        "fields": ["Employer Name", "Employee Name", "Job Title", "Company Name", "Salary", "Payment Period", "Notice Period", "Effective Date"],
        "content": """This Employment Agreement ('Agreement') is entered into on {Effective Date} between:

{Employer Name} (Employer) and {Employee Name} (Employee).

1. Position and Responsibilities:
   The Employee will work as {Job Title} at {Company Name}.

2. Compensation:
   The Employee will be paid {Salary} per {Payment Period}.

3. Termination:
   This Agreement may be terminated by either party with {Notice Period} notice.

"""
    }
}

# Function to create a PDF document
def create_pdf(filepath, document_data, filled_content):
    c = canvas.Canvas(filepath, pagesize=letter)
    
    # Set margins and content width
    margin = 50
    content_width = 500
    y_position = 750  # Starting position for text

    # Title (Centered)
    c.setFont("Helvetica-Bold", 16)
    c.drawCentredString(300, y_position, document_data["title"])
    y_position -= 30  # Space below title

    # Body text settings
    c.setFont("Helvetica", 12)
    
    # Draw text with proper wrapping and margins
    for line in filled_content.split("\n"):
        wrapped_lines = simpleSplit(line.strip(), "Helvetica", 12, content_width)
        for wrapped_line in wrapped_lines:
            c.drawString(margin, y_position, wrapped_line)
            y_position -= 20  # Space between lines
            if y_position < 50:  # New page if needed
                c.showPage()
                c.setFont("Helvetica", 12)
                y_position = 750

    # Draw Signature Lines (Properly Aligned)
    y_position -= 40
    c.line(margin, y_position, margin + 200, y_position)  # Disclosing Party
    c.line(margin + 300, y_position, margin + 500, y_position)  # Date

    c.drawString(margin, y_position - 15, "Signed: (Disclosing Party)")
    c.drawString(margin + 300, y_position - 15, "Date:")

    y_position -= 40
    c.line(margin, y_position, margin + 200, y_position)  # Receiving Party
    c.line(margin + 300, y_position, margin + 500, y_position)  # Date

    c.drawString(margin, y_position - 15, "Signed: (Receiving Party)")
    c.drawString(margin + 300, y_position - 15, "Date:")

    c.save()


# Function to create a well-formatted Word document
def create_word(filepath, document_data, filled_content):
    doc = Document()

    # Add Title (Centered, Bold, Larger Font)
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run(document_data["title"])
    title_run.bold = True
    title_run.font.size = Pt(16)

    # Add spacing after title
    doc.add_paragraph("\n")

    # Add content with justified alignment
    for line in filled_content.split("\n"):
        para = doc.add_paragraph(line.strip())
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        para.space_after = Pt(10)

    # Add signature fields properly aligned
    doc.add_paragraph("\n")
    table = doc.add_table(rows=2, cols=2)
    table.autofit = True

    # Signature lines
    table.cell(0, 0).text = "Signed: ______________________ (Disclosing Party)"
    table.cell(0, 1).text = "Date: ______________________"

    table.cell(1, 0).text = "Signed: ______________________ (Receiving Party)"
    table.cell(1, 1).text = "Date: ______________________"

    # Save the document
    doc.save(filepath)

# Function to generate the document
def generate_document():
    doc_type = document_var.get()
    save_format = format_var.get()
    print(doc_type, save_format)
    if not doc_type or doc_type not in document_templates:

        messagebox.showerror("Error", "Please select a valid document type.")
        return
    
        # Check if format is selected
    if save_format not in ["PDF", "Word", "Both"]:
        messagebox.showerror("Error", "Please select a valid file type.")
        return
    document_data = document_templates[doc_type]

    
    user_inputs = {}
    for field in document_data["fields"]:
        value = input_fields[field].get().strip()
        if not value:
            messagebox.showerror("Error", f"Please fill in {field}.")
            return
        user_inputs[field] = value

    filled_content = document_data["content"].format(**user_inputs)

    # Ask for save location
    filetypes = [("PDF Files", "*.pdf"), ("Word Files", "*.docx")]
    save_path = filedialog.asksaveasfilename(defaultextension=".pdf" if save_format == "PDF" else ".docx",
                                             filetypes=filetypes,
                                             title="Choose Save Location")
    if not save_path:
        return

      # Save in selected format
    if save_format == "PDF":
        create_pdf(save_path, document_data, filled_content)
    elif save_format == "Word":
        create_word(save_path, document_data, filled_content)
    elif save_format == "Both":
        create_pdf(save_path.replace(".docx", ".pdf"), document_data, filled_content)
        create_word(save_path.replace(".pdf", ".docx"), document_data, filled_content)

    messagebox.showinfo("Success", f"{document_data['title']} saved successfully!")


# Function to update input fields dynamically
def update_fields(*args):
    selected_doc = document_var.get()
    
    for widget in field_frame.winfo_children():
        widget.destroy()  # Clear old fields

    if selected_doc in document_templates:
        document_data = document_templates[selected_doc]
        tk.Label(field_frame, text="Enter Details:", font=("Arial", 12, "bold")).pack(anchor="w")

        for field in document_data["fields"]:
            tk.Label(field_frame, text=field + ":", font=("Arial", 10)).pack(anchor="w")
            entry = tk.Entry(field_frame, width=40)
            entry.pack()
            input_fields[field] = entry

# Initialize Tkinter window
root = tk.Tk()
root.title("Legal Document Generator")
root.geometry("500x600")
#root.iconbitmap("C:\Users\prana\Downloads\LegalAssist\LegalAssist\legal-document.ico")
root.configure(padx=20, pady=20)

#toggling between dark and light mode
def toggle_theme():
    current_theme = style.theme_use()
    new_theme = "darkly" if current_theme == "cosmo" else "cosmo"
    style.theme_use(new_theme)
    toggle_btn.config(text="🌙 Dark Mode" if new_theme == "cosmo" else "☀️ Light Mode")

style = ttk.Style()
#Header
ttk.Label(root, text="📜 Legal Document Generator", font=("Arial", 16, "bold"), bootstyle="primary").pack(pady=10)


# Dropdown menu for document selection
ttk.Label(root, text="Select a Legal Document:", font=("Arial", 12, "bold")).pack(pady=5)
document_var = ttk.StringVar()
document_var.trace("w", update_fields)

document_menu = ttk.Combobox(root, textvariable=document_var, values=list(document_templates.keys()), font=("Arial", 12))
document_menu.pack(pady=5)

# Dynamic input fields frame
field_frame = ttk.Frame(root)
field_frame.pack(pady=10)

input_fields = {}

# Variables for document type and format
doc_type_var = StringVar()
format_var = StringVar(value="PDF")

#buttons for selecting format
ttk.Label(root, text="Select Format:", font=("Arial", 12, "bold")).pack(pady=5)
pdf_radio = ttk.Radiobutton(root, text="PDF", variable=format_var, value="PDF")
pdf_radio.pack()
word_radio = ttk.Radiobutton(root, text="Word", variable=format_var, value="Word")
word_radio.pack()
both_radio = ttk.Radiobutton(root, text="Both", variable=format_var, value="Both")
both_radio.pack()


# Generate button
generate_btn = ttk.Button(root, text="Generate Document", command=generate_document, bootstyle="success")
generate_btn.pack(pady=20)

#toggle_button
toggle_btn = ttk.Button(root, text="🌙 Niggah Mode", command=toggle_theme, bootstyle="primary")
toggle_btn.pack(pady=10)

toggle_theme()
root.mainloop()
